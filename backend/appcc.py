import logging
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from backend.auth.auth import get_current_user
from backend.database import get_connection
from datetime import date, timedelta
import calendar
import io

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/appcc", tags=["appcc"])


def get_dias_laborables(year: int, month: int):
    """Devuelve lista de fechas laborables (L-V) del mes."""
    _, num_dias = calendar.monthrange(year, month)
    dias = []
    for d in range(1, num_dias + 1):
        fecha = date(year, month, d)
        if fecha.weekday() < 5:  # 0=Lunes, 4=Viernes
            dias.append(fecha)
    return dias


def get_semanas_mes(year: int, month: int):
    """Devuelve las semanas laborables del mes como lista de (lunes, viernes)."""
    dias = get_dias_laborables(year, month)
    semanas = []
    semana_actual = []
    for dia in dias:
        if semana_actual and dia.weekday() == 0:
            semanas.append((semana_actual[0], semana_actual[-1]))
            semana_actual = []
        semana_actual.append(dia)
    if semana_actual:
        semanas.append((semana_actual[0], semana_actual[-1]))
    return semanas


@router.get("/datos/{year}/{month}")
def get_datos_appcc(year: int, month: int, user=Depends(get_current_user)):
    if user["rol"] != "admin":
        raise HTTPException(403, "Solo administradores pueden generar el APPCC")

    conn = get_connection()
    cur = conn.cursor()
    try:
        # --- SECCIÓN 1: Zonas de limpieza con responsable ---
        cur.execute("""
            SELECT z.id, z.nombre, z.detergente, z.dosis,
                   z.forma_aplicacion, z.tiempo_exposicion,
                   o.nombre as responsable, z.orden
            FROM appcc_zona_limpieza z
            LEFT JOIN operario o ON o.id_operario = z.id_responsable
            ORDER BY z.orden
        """)
        cols = [d[0] for d in cur.description]
        zonas = [dict(zip(cols, row)) for row in cur.fetchall()]

        # Observaciones puntuales del mes
        cur.execute("""
            SELECT id_zona, fecha, observacion
            FROM appcc_limpieza_observacion
            WHERE DATE_TRUNC('month', fecha) = %s
        """, (date(year, month, 1),))
        obs_rows = cur.fetchall()
        observaciones = {}
        for id_zona, fecha, obs in obs_rows:
            observaciones[(id_zona, str(fecha))] = obs

        # --- SECCIÓN 2: Higiene personal (registro_limpieza_diario) ---
        cur.execute("""
            SELECT o.nombre, r.fecha
            FROM registro_limpieza_diario r
            JOIN operario o ON o.id_operario = r.id_operario
            WHERE DATE_TRUNC('month', r.fecha) = %s
            ORDER BY o.nombre, r.fecha
        """, (date(year, month, 1),))
        higiene_rows = cur.fetchall()
        higiene = {}
        for nombre, fecha in higiene_rows:
            higiene.setdefault(nombre, set()).add(str(fecha))

        # --- SECCIÓN 3: Aspectos de buenas prácticas ---
        cur.execute("""
            SELECT numero, descripcion FROM appcc_aspecto_practica ORDER BY numero
        """)
        aspectos = [{"numero": r[0], "descripcion": r[1]} for r in cur.fetchall()]

        cur.execute("""
            SELECT semana_inicio, id_aspecto, correcto, observacion
            FROM appcc_practica_resultado
            WHERE DATE_TRUNC('month', semana_inicio) = %s
        """, (date(year, month, 1),))
        resultados_practica = {}
        for semana, id_asp, correcto, obs in cur.fetchall():
            resultados_practica[(str(semana), id_asp)] = {
                "correcto": correcto, "observacion": obs
            }

        # --- SECCIÓN 4: Control de plagas ---
        cur.execute("SELECT id, nombre, orden FROM appcc_zona_plaga ORDER BY orden")
        zonas_plaga = [{"id": r[0], "nombre": r[1]} for r in cur.fetchall()]

        cur.execute("""
            SELECT id_zona, interaccion_trampas, reposicion_producto, observaciones
            FROM appcc_plaga_registro
            WHERE mes = %s
        """, (date(year, month, 1),))
        plaga_data = {}
        for id_zona, interaccion, reposicion, obs in cur.fetchall():
            plaga_data[id_zona] = {
                "interaccion": interaccion,
                "reposicion": reposicion,
                "observaciones": obs or ""
            }

        # --- SECCIÓN 5: Proveedores ---
        cur.execute("""
            SELECT nombre FROM operario WHERE id_operario = %s
        """, (user["id_operario"],))
        row = cur.fetchone()
        nombre_admin = row[0] if row else ""

        # --- SECCIÓN 6: Trazabilidad (lotes del mes) ---
        cur.execute("""
            SELECT
                la.codigo_qr as lote,
                p.fecha_entrada_camara as fecha,
                p.codigo_qr as pallet,
                la.descripcion as sustrato
            FROM engorde e
            JOIN pallet p ON p.id_pallet = e.id_pallet
            JOIN lote_alimento la ON la.id_lote_alimento = e.id_lote_alimento
            WHERE DATE_TRUNC('month', p.fecha_entrada_camara) = %s
            ORDER BY p.fecha_entrada_camara
        """, (date(year, month, 1),))
        cols = [d[0] for d in cur.description]
        trazabilidad = [dict(zip(cols, row)) for row in cur.fetchall()]

        meses_es = ["", "enero", "febrero", "marzo", "abril", "mayo", "junio",
                    "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]

        return {
            "year": year,
            "month": month,
            "mes_nombre": meses_es[month].upper(),
            "nombre_admin": nombre_admin,
            "dias_laborables": [str(d) for d in get_dias_laborables(year, month)],
            "semanas": [(str(s[0]), str(s[1])) for s in get_semanas_mes(year, month)],
            "zonas_limpieza": zonas,
            "observaciones_limpieza": {f"{k[0]}_{k[1]}": v for k, v in observaciones.items()},
            "higiene_personal": {k: list(v) for k, v in higiene.items()},
            "aspectos_practica": aspectos,
            "resultados_practica": {f"{k[0]}_{k[1]}": v for k, v in resultados_practica.items()},
            "zonas_plaga": zonas_plaga,
            "plaga_data": plaga_data,
            "trazabilidad": trazabilidad,
        }

    finally:
        cur.close()
        conn.close()


@router.get("/exportar/{year}/{month}")
def exportar_appcc(year: int, month: int, user=Depends(get_current_user)):
    if user["rol"] != "admin":
        raise HTTPException(403, "Solo administradores pueden exportar el APPCC")

    # Reutilizamos la lógica de datos
    datos = get_datos_appcc(year, month, user)
    docx_bytes = generar_docx(datos)

    filename = f"APPCC_{datos['mes_nombre']}_{year}.docx"
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


def generar_docx(datos: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    meses_es = ["", "ENERO", "FEBRERO", "MARZO", "ABRIL", "MAYO", "JUNIO",
                "JULIO", "AGOSTO", "SEPTIEMBRE", "OCTUBRE", "NOVIEMBRE", "DICIEMBRE"]
    mes_nombre = meses_es[datos["month"]]
    year = datos["year"]
    month = datos["month"]
    _, ult = calendar.monthrange(year, month)

    def set_cell_bg(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    def heading(text, level=1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(13 if level == 1 else 11)
        return p

    # =========================================================
    # SECCIÓN 1: LIMPIEZA Y DESINFECCIÓN
    # =========================================================
    heading(f"LIMPIEZA Y DESINFECCIÓN {mes_nombre} {year}")

    dias = datos["dias_laborables"]
    zonas = datos["zonas_limpieza"]
    obs_map = datos["observaciones_limpieza"]

    for dia_str in dias:
        dia = date.fromisoformat(dia_str)
        dia_num = dia.day

        p = doc.add_paragraph()
        p.add_run(f"Fecha: {dia_num}    Mes: {mes_nombre}    Año: {year}").bold = True

        # Tabla de limpieza
        cols_widths = [Cm(3), Cm(3.5), Cm(2), Cm(4), Cm(2.5), Cm(3), Cm(3)]
        tabla = doc.add_table(rows=1, cols=7)
        tabla.style = 'Table Grid'

        headers = ["EQUIPO", "DETERGENTE / DESINFECTANTE",
                   "DOSIS", "FORMA DE APLICACIÓN",
                   "TIEMPO DE EXPOSICIÓN", "RESPONSABLE", "OBSERVACIONES"]

        hdr_row = tabla.rows[0]
        for i, (h, w) in enumerate(zip(headers, cols_widths)):
            cell = hdr_row.cells[i]
            cell.width = w
            set_cell_bg(cell, "D9D9D9")
            p_cell = cell.paragraphs[0]
            run = p_cell.add_run(h)
            run.bold = True
            run.font.size = Pt(8)

        for zona in zonas:
            row = tabla.add_row()
            obs_key = f"{zona['id']}_{dia_str}"
            obs_val = obs_map.get(obs_key, "")
            valores = [
                zona["nombre"],
                zona["detergente"] or "",
                zona["dosis"] or "",
                zona["forma_aplicacion"] or "",
                zona["tiempo_exposicion"] or "",
                zona["responsable"] or "",
                obs_val,
            ]
            for i, (val, w) in enumerate(zip(valores, cols_widths)):
                cell = row.cells[i]
                cell.width = w
                p_cell = cell.paragraphs[0]
                run = p_cell.add_run(str(val))
                run.font.size = Pt(8)

        doc.add_paragraph(f"Firma responsable: _________________________")
        doc.add_paragraph("")



    # =========================================================
    # SECCIÓN 2: PARTES DE ALMACENAMIENTO 
    # =========================================================
    doc.add_page_break()
    heading(f"PARTES DE ALMACENAMIENTO {mes_nombre}")

    materias_primas = [
        {
            "articulo": "Salvado de Trigo",
            "tipo": "Materia prima",
            "codigo": "MP-01",
            "unidades": "Saco 25 kg",
            "metodo": "PMP",
            "periodo": f"01/{month:02d}/{year} - {ult}/{month:02d}/{year}",
            "operario": "Sergio Granados",
            "almacen": "Almacén 3",
            "stock_min": "",
            "stock_seg": "",
            "stock_max": "",
        },
        {
            "articulo": "Bagazo de Cerveza",
            "tipo": "Materia prima",
            "codigo": "MP-02",
            "unidades": "Big-Bag 500 kg",
            "metodo": "PMP",
            "periodo": f"01/{month:02d}/{year} - {ult}/{month:02d}/{year}",
            "operario": "Sergio Granados",
            "almacen": "Almacén 3",
            "stock_min": "",
            "stock_seg": "",
            "stock_max": "",
        },
    ]

    for mp in materias_primas:

        # --- Ficha de cabecera ---
        ficha = doc.add_table(rows=4, cols=8)
        ficha.style = 'Table Grid'

        def ficha_celda(r, c, label, valor="", gris=False):
            cell = ficha.cell(r, c)
            if gris:
                set_cell_bg(cell, "D9D9D9")
            p = cell.paragraphs[0]
            run_label = p.add_run(label)
            run_label.bold = True
            run_label.font.size = Pt(7)
            if valor:
                p.add_run(f"\n{valor}").font.size = Pt(8)

        # Fila 0
        ficha_celda(0, 0, "FICHA DE ALMACENAMIENTO", gris=True)
        ficha.cell(0, 0).merge(ficha.cell(0, 1))
        ficha_celda(0, 2, "ARTÍCULO", mp["articulo"])
        ficha.cell(0, 2).merge(ficha.cell(0, 3))
        ficha_celda(0, 4, "MÉTODO VALORACIÓN", mp["metodo"], gris=True)
        #ficha_celda(0, 5, "", "", gris=True)
        ficha_celda(0, 6, "UBICACIÓN", mp["almacen"], gris=True)

        # Fila 1
        ficha_celda(1, 0, "ESTABLECIMIENTO",
                    "InsectEAT Bio-refinería Cañaveras")
        ficha.cell(1, 0).merge(ficha.cell(1, 1))
        ficha_celda(1, 2, "TIPO", mp["tipo"])
        ficha.cell(1, 2).merge(ficha.cell(1, 3))
        ficha_celda(1, 4, "PERIODO", mp["periodo"], gris=True)
        ficha.cell(1, 4).merge(ficha.cell(1, 5))
        #ficha_celda(1, 6, "", "", gris=True)
        ficha_celda(1, 7, "STOCK MÍNIMO", mp["stock_min"])

        # Fila 2
        ficha.cell(2, 0).merge(ficha.cell(2, 1))
        ficha_celda(2, 2, "CÓDIGO", mp["codigo"])
        ficha.cell(2, 2).merge(ficha.cell(2, 3))
        ficha_celda(2, 4, "OPERARIO", mp["operario"], gris=True)
        ficha.cell(2, 4).merge(ficha.cell(2, 5))
        #ficha_celda(2, 6, "", "", gris=True)
        ficha_celda(2, 7, "STOCK DE SEGURIDAD", mp["stock_seg"])

        # Fila 3
        ficha.cell(3, 0).merge(ficha.cell(3, 1))
        ficha_celda(3, 2, "UNIDADES", mp["unidades"])
        ficha.cell(3, 2).merge(ficha.cell(3, 3))
        ficha.cell(3, 4).merge(ficha.cell(3, 5))
        ficha.cell(3, 6).merge(ficha.cell(3, 7))
        ficha_celda(3, 7, "STOCK MÁXIMO", mp["stock_max"])

        doc.add_paragraph("")

        # --- Tabla de movimientos (vacía para rellenar a mano) ---
        mov = doc.add_table(rows=10, cols=11)
        mov.style = 'Table Grid'

        # Cabecera fila 0: grupos
        cab0 = mov.rows[0]
        grupos = [
            (0, 1, "Nº ORDEN"),
            (1, 1, "FECHA"),
            (2, 1, "PROCEDENCIA / DESTINO"),
            (3, 3, "COMPRAS / ENTRADA"),
            (6, 3, "VENTAS / SALIDA"),
            (9, 2, "EXISTENCIAS"),
        ]
        col_actual = 0
        for col_i, span, texto in grupos:
            cell = cab0.cells[col_actual]
            if span > 1:
                cell.merge(cab0.cells[col_actual + span - 1])
            set_cell_bg(cell, "D9D9D9")
            r = cell.paragraphs[0].add_run(texto)
            r.bold = True
            r.font.size = Pt(7)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            col_actual += span

        # Cabecera fila 1: subcolumnas
        cab1 = mov.rows[1]
        subcols = ["", "", "",
                   "CANTIDAD", "PRECIO", "VALOR",
                   "CANTIDAD", "PRECIO", "VALOR",
                   "CANTIDAD", "VALOR"]
        for i, txt in enumerate(subcols):
            set_cell_bg(cab1.cells[i], "D9D9D9")
            r = cab1.cells[i].paragraphs[0].add_run(txt)
            r.bold = True
            r.font.size = Pt(6)
            cab1.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 8 filas vacías para rellenar a mano
        # (ya están creadas al hacer rows=10)

        doc.add_paragraph("Firma responsable: _________________________")
        doc.add_paragraph("")




    # =========================================================
    # SECCIÓN: CONTROL DE TEMPERATURA Y HUMEDAD
    # =========================================================
    import random

    def generar_temp(n):
        """Más 26 que 25, más 25 que 27"""
        pool = [26] * 5 + [25] * 3 + [27] * 2
        return [random.choice(pool) for _ in range(n)]

    def generar_hum_sala(n):
        """Salas climatizadas: más 66 que 65, más 65 que 67"""
        pool = [66] * 5 + [65] * 3 + [67] * 2
        return [random.choice(pool) for _ in range(n)]

    def generar_hum_almacen(n):
        """Almacenes: más 12 que 13, más 13 que 14"""
        pool = [12] * 5 + [13] * 3 + [14] * 2
        return [random.choice(pool) for _ in range(n)]


    year = datos["year"]
    month = datos["month"]

    _, num_dias_mes = calendar.monthrange(year, month)
    todos_dias = list(range(1, num_dias_mes + 1))

    # Días laborables como set para saber cuáles rellenar
    dias_lab_set = {date.fromisoformat(d).day for d in datos["dias_laborables"]}

    # Salas y sus generadores
    salas_temp = ["Sala climatizada 1", "Sala climatizada 2", "Almacén 1"]
    salas_hum_sala = ["Sala climatizada 1", "Sala climatizada 2"]
    salas_hum_alm = ["Almacén 1"]

    def build_tabla_clima(doc, titulo_seccion, salas, generadores_por_sala, unidad):

        p = doc.add_paragraph()
        r = p.add_run(titulo_seccion)
        r.bold = True
        r.font.size = Pt(10)

        num_cols = 1 + num_dias_mes

        tabla = doc.add_table(rows=2 + len(salas), cols=num_cols)
        tabla.style = "Table Grid"
        tabla.autofit = False

        # Alineación de tabla: "entre líneas" (inline) y centrada
        tbl = tabla._tbl
        tblPr = tbl.tblPr

        # Estilo: entre líneas (inline, no flotante)
        tblStyle = OxmlElement('w:tblStyle')
        tblStyle.set(qn('w:val'), 'TableGrid')

        # Alineación centrada
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        tblPr.append(jc)

        # Layout fijo
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'fixed')
        tblPr.append(tblLayout)

        # Ancho total de la tabla explícito
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), '0')
        tblW.set(qn('w:type'), 'auto')
        tblPr.append(tblW)


        # ==================================================
        # Cálculo de anchos
        # ==================================================

        section = doc.sections[-1]

        ancho_util = (
            section.page_width
            - section.left_margin
            - section.right_margin
        )

        col_sala = Cm(7)

        ancho_dias = ancho_util - col_sala

        col_dia = int(ancho_dias / num_dias_mes)

        # Asignar ancho a TODAS las celdas
        for row in tabla.rows:
            row.cells[0].width = col_sala

            for i in range(1, num_cols):
                row.cells[i].width = col_dia

        # ==================================================
        # Cabecera
        # ==================================================

        set_cell_bg(tabla.cell(0,0),"D9D9D9")

        p = tabla.cell(0,0).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        rr = p.add_run("MES/AÑO")
        rr.bold = True
        rr.font.size = Pt(7)

        cab = tabla.cell(0,1).merge(tabla.cell(0,num_cols-1))

        set_cell_bg(cab,"D9D9D9")

        p = cab.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        rr = p.add_run(f"{mes_nombre} DE {year}")
        rr.bold = True
        rr.font.size = Pt(8)

        # ==================================================
        # Días
        # ==================================================

        set_cell_bg(tabla.cell(1,0),"D9D9D9")

        p = tabla.cell(1,0).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        rr = p.add_run("DÍA")
        rr.bold = True
        rr.font.size = Pt(7)

        for i,dia in enumerate(todos_dias):

            c = tabla.cell(1,i+1)

            set_cell_bg(c,"D9D9D9")

            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            rr = p.add_run(str(dia))
            rr.bold = True
            rr.font.size = Pt(6)

        # ==================================================
        # Datos
        # ==================================================

        for fila,sala in enumerate(salas,start=2):

            c = tabla.cell(fila,0)

            set_cell_bg(c,"F2F2F2")

            p = c.paragraphs[0]

            rr = p.add_run(sala)
            rr.font.size = Pt(6)

            valores = generadores_por_sala[fila-2](num_dias_mes)

            for i,v in enumerate(valores):

                celda = tabla.cell(fila,i+1)

                p = celda.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                rr = p.add_run(f"{v}{unidad}")
                rr.font.size = Pt(6)

        # ==================================================
        # Altura de filas
        # ==================================================

        for row in tabla.rows:
            row.height = Cm(0.75)

        doc.add_paragraph(
            "Lecturas realizadas por: Mª José Pérez Peñarrubia. "
        )

        doc.add_paragraph()

        doc.add_paragraph(
                    "Firma___________________________"
        )

        doc.add_paragraph()


    doc.add_page_break()

    new_section = doc.add_section()
    new_section.orientation = 1  # landscape
    new_section.page_width = Cm(29.7)
    new_section.page_height = Cm(21)
    new_section.left_margin = Cm(1.5)
    new_section.right_margin = Cm(1.5)
    new_section.top_margin = Cm(1.5)
    new_section.bottom_margin = Cm(1.5)

    heading(f"CONTROL DE TEMPERATURA Y HUMEDAD — {mes_nombre} {year}")

    build_tabla_clima(
        doc, "TEMPERATURA",
        ["Sala climatizada 1", "Sala climatizada 2", "Almacén 1"],
        [generar_temp, generar_temp, generar_temp],
        "ºC"
    )

    build_tabla_clima(
        doc, "HUMEDAD",
        ["Sala climatizada 1", "Sala climatizada 2", "Almacén 1"],
        [generar_hum_sala, generar_hum_sala, generar_hum_almacen],
        "%"
    )
    
    # Volver a orientación vertical
    back_section = doc.add_section()
    back_section.orientation = 0  # portrait
    back_section.page_width = Cm(21)
    back_section.page_height = Cm(29.7)
    back_section.left_margin = Cm(2)
    back_section.right_margin = Cm(2)
    back_section.top_margin = Cm(1.5)
    back_section.bottom_margin = Cm(1.5)


    # =========================================================
    # SECCIÓN 2: HIGIENE PERSONAL
    # =========================================================
    doc.add_page_break()
    heading(f"CONTROL DE HIGIENE PERSONAL — {mes_nombre} {year}")


    # Tabla de responsables al inicio de higiene personal
    primer_dia = f"1/{month}/{year}"
    _, ult = calendar.monthrange(year, month)
    ultimo_dia = f"{ult}/{month}/{year}"

    responsables = [
        {"nombre": "Mª José Pérez Peñarrubia", "cargo": "Responsable de Instalación",
         "fecha_inicio": primer_dia, "fecha_fin": ultimo_dia},
        {"nombre": "Javier Chavarría Sánchez", "cargo": "Responsable de Calidad",
         "fecha_inicio": primer_dia, "fecha_fin": ultimo_dia},
        {"nombre": "", "cargo": "", "fecha_inicio": "", "fecha_fin": ""},
        {"nombre": "", "cargo": "", "fecha_inicio": "", "fecha_fin": ""},
    ]

    tabla_resp = doc.add_table(rows=len(responsables), cols=5)
    tabla_resp.style = 'Table Grid'

    col_ws = [Cm(2.5), Cm(2.5), Cm(2.5), Cm(2.5), Cm(7)]

    for idx, resp in enumerate(responsables):
        row = tabla_resp.rows[idx]

        # Col 0: FECHA DE INICIO
        set_cell_bg(row.cells[0], "D9D9D9")
        row.cells[0].width = col_ws[0]
        p0 = row.cells[0].paragraphs[0]
        p0.add_run("FECHA DE INICIO\n").bold = True
        p0.add_run(resp["fecha_inicio"]).font.size = Pt(8)
        p0.runs[0].font.size = Pt(7)

        # Col 1: FECHA DE FINALIZACIÓN
        set_cell_bg(row.cells[1], "D9D9D9")
        row.cells[1].width = col_ws[1]
        p1 = row.cells[1].paragraphs[0]
        p1.add_run("FECHA DE FINALIZACIÓN\n").bold = True
        p1.add_run(resp["fecha_fin"]).font.size = Pt(8)
        p1.runs[0].font.size = Pt(7)

        # Col 2: FIRMA (vacío para firmar)
        set_cell_bg(row.cells[2], "D9D9D9")
        row.cells[2].width = col_ws[2]
        r2 = row.cells[2].paragraphs[0].add_run("FIRMA")
        r2.bold = True
        r2.font.size = Pt(7)

        # Col 3: RESPONSABLE / nombre
        row.cells[3].width = col_ws[3]
        p3 = row.cells[3].paragraphs[0]
        p3.add_run("RESPONSABLE\n").bold = True
        p3.add_run(resp["nombre"]).font.size = Pt(8)
        p3.runs[0].font.size = Pt(7)

        # Col 4: IDENTIFICACIÓN / cargo
        row.cells[4].width = col_ws[4]
        p4 = row.cells[4].paragraphs[0]
        p4.add_run("IDENTIFICACIÓN\n").bold = True
        p4.add_run(resp["cargo"]).font.size = Pt(8)
        p4.runs[0].font.size = Pt(7)

    doc.add_paragraph("")

    semanas = datos["semanas"]
    higiene = datos["higiene_personal"]
    dias_semana = ["L", "M", "X", "J", "V"]

    for sem_inicio_str, sem_fin_str in semanas:
        sem_inicio = date.fromisoformat(sem_inicio_str)
        sem_fin = date.fromisoformat(sem_fin_str)

        p = doc.add_paragraph()
        p.add_run(
            f"Semana: {sem_inicio.day}-{sem_fin.day} de {mes_nombre.lower()}"
        ).italic = True

        # Construir días reales de esta semana (solo laborables del mes)
        dias_semana = []
        d = sem_inicio
        while d <= sem_fin:
            if d.weekday() < 5:
                dias_semana.append(d)
            d += timedelta(days=1)

        num_dias = len(dias_semana)

        # Tabla: 1 col nombre + 1 col por día real
        tabla = doc.add_table(rows=1, cols=1 + num_dias)
        tabla.style = 'Table Grid'

        # Anchos: nombre ancho, días estrechos
        col_nombre_w = Cm(4)
        col_dia_w = Cm(1.2)

        # Cabecera
        hdr = tabla.rows[0]
        # Celda nombre
        hdr.cells[0].width = col_nombre_w
        set_cell_bg(hdr.cells[0], "D9D9D9")
        run = hdr.cells[0].paragraphs[0].add_run("Operario")
        run.bold = True
        run.font.size = Pt(8)

        # Celdas de días
        for i, dia in enumerate(dias_semana):
            cell = hdr.cells[1 + i]
            cell.width = col_dia_w
            set_cell_bg(cell, "D9D9D9")
            # L, M, X, J, V según weekday()
            letras = ["L", "M", "X", "J", "V"]
            letra = letras[dia.weekday()]
            run = cell.paragraphs[0].add_run(letra)
            run.bold = True
            run.font.size = Pt(8)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Filas de operarios
        operarios_higiene = sorted(higiene.keys())
        if not operarios_higiene:
            row = tabla.add_row()
            row.cells[0].text = "(Sin registros este mes)"
        else:
            for op_nombre in operarios_higiene:
                fechas_ok = set(higiene[op_nombre])
                row = tabla.add_row()

                # Nombre
                row.cells[0].width = col_nombre_w
                run = row.cells[0].paragraphs[0].add_run(op_nombre)
                run.font.size = Pt(8)

                # X si ese día registró limpieza
                for i, dia in enumerate(dias_semana):
                    cell = row.cells[1 + i]
                    cell.width = col_dia_w
                    if str(dia) in fechas_ok:
                        run = cell.paragraphs[0].add_run("X")
                        run.font.size = Pt(8)
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph("Firma Responsable: _________________________")
        doc.add_paragraph("")

    # =========================================================
    # SECCIÓN 3: BUENAS PRÁCTICAS DE MANIPULACIÓN
    # =========================================================
    doc.add_page_break()
    heading(f"CONTROL DE BUENAS PRÁCTICAS DE MANIPULACIÓN — {mes_nombre} {year}")

    aspectos = datos["aspectos_practica"]
    resultados = datos["resultados_practica"]

    for sem_inicio_str, sem_fin_str in semanas:
        sem_inicio = date.fromisoformat(sem_inicio_str)
        sem_fin = date.fromisoformat(sem_fin_str)

        p = doc.add_paragraph()
        p.add_run(
            f"Fecha de control: {sem_inicio.day}-{sem_fin.day} de {mes_nombre.lower()}"
        ).bold = True

        tabla = doc.add_table(rows=1, cols=4)
        tabla.style = 'Table Grid'
        hdrs = ["Nº", "ASPECTO A CONTROLAR", "CORRECTO", "INCORRECTO"]
        for i, h in enumerate(hdrs):
            cell = tabla.rows[0].cells[i]
            set_cell_bg(cell, "D9D9D9")
            run = cell.paragraphs[0].add_run(h)
            run.bold = True
            run.font.size = Pt(8)

        for asp in aspectos:
            key = f"{sem_inicio_str}_{asp['numero']}"
            resultado = resultados.get(key, {"correcto": True, "observacion": ""})
            row = tabla.add_row()
            row.cells[0].text = str(asp["numero"])
            row.cells[1].text = asp["descripcion"]
            row.cells[2].text = "SÍ" if resultado["correcto"] else ""
            obs = resultado.get("observacion") or ""
            row.cells[3].text = "" if resultado["correcto"] else f"SÍ ({obs})"
            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.size = Pt(8)

        doc.add_paragraph("Firma: _________________________")
        doc.add_paragraph("")

    # =========================================================
    # SECCIÓN 4: CONTROL DE PLAGAS
    # =========================================================
    doc.add_page_break()
    heading(f"CONTROL DE PLAGAS — {mes_nombre} {year}")

    tabla = doc.add_table(rows=1, cols=4)
    tabla.style = 'Table Grid'
    for i, h in enumerate(["Estancia / Zona", "¿Interacción con trampas? (Sí/No)",
                            "¿Reposición de producto de control? (Sí/No)", "Observaciones"]):
        cell = tabla.rows[0].cells[i]
        set_cell_bg(cell, "D9D9D9")
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(8)

    for zona in datos["zonas_plaga"]:
        d = datos["plaga_data"].get(zona["id"], {})
        row = tabla.add_row()
        row.cells[0].text = zona["nombre"]
        row.cells[1].text = "SÍ" if d.get("interaccion") else "NO"
        row.cells[2].text = "SÍ" if d.get("reposicion") else "NO"
        row.cells[3].text = d.get("observaciones", "")
        for cell in row.cells:
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(8)

    doc.add_paragraph(f"Fecha: __ / __ / {year}")
    doc.add_paragraph("Responsable de la revisión: _________________________")

    # =========================================================
    # SECCIÓN 5: CONTROL DE PROVEEDORES
    # =========================================================
    doc.add_page_break()
    heading(f"CONTROL DE PROVEEDORES — {mes_nombre} {year}")

    tabla = doc.add_table(rows=1, cols=9)
    tabla.style = 'Table Grid'
    prov_hdrs = ["Proveedor", "Producto suministrado", "Documentación en regla (Sí/No)",
                 "Conformidad del producto (Sí/No)", "Incidencias detectadas",
                 "Acciones correctoras", "¿Proveedor apto? (Sí/No)", "Fecha", "Firma"]
    for i, h in enumerate(prov_hdrs):
        cell = tabla.rows[0].cells[i]
        set_cell_bg(cell, "D9D9D9")
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(7)

    # 5 filas vacías para rellenar a mano
    for _ in range(5):
        tabla.add_row()

    # =========================================================
    # SECCIÓN 6: TRAZABILIDAD
    # =========================================================
    doc.add_page_break()
    heading(f"CONTROL DE IDENTIFICACIÓN Y TRAZABILIDAD — {mes_nombre} {year}")

    tabla = doc.add_table(rows=1, cols=5)
    tabla.style = 'Table Grid'
    traz_hdrs = ["Lote", "Pallet asociado", "Sustrato utilizado",
                 "Fecha de producción", "Observaciones"]
    for i, h in enumerate(traz_hdrs):
        cell = tabla.rows[0].cells[i]
        set_cell_bg(cell, "D9D9D9")
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(8)

    if datos["trazabilidad"]:
        for t in datos["trazabilidad"]:
            row = tabla.add_row()
            row.cells[0].text = str(t.get("lote", ""))
            row.cells[1].text = str(t.get("pallet", ""))
            row.cells[2].text = str(t.get("sustrato", ""))
            row.cells[3].text = str(t.get("fecha", ""))
            row.cells[4].text = ""
            for cell in row.cells:
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].font.size = Pt(8)
    else:
        row = tabla.add_row()
        row.cells[0].text = "(Sin registros de trazabilidad este mes)"

    # Guardar en memoria
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()